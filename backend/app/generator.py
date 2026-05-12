import os
import json
import time
import logging
from docx import Document
from docx.shared import Pt
from .llm import generate_text

logger = logging.getLogger("docgen.generator")

BASE_DIR = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(BASE_DIR, "template")
TEMPLATE_FILE = os.path.join(TEMPLATE_DIR, "hld_template.docx")
PROMPT_FILE = os.path.join(TEMPLATE_DIR, "prompt.json")


def ensure_template():
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    if not os.path.exists(PROMPT_FILE):
        default = [
            {"name": "Overview", "type": "text", "instruction": "Provide a concise overview of the application."},
            {"name": "Architecture", "type": "text", "instruction": "High-level architecture description and components."},
            {"name": "Components", "type": "table", "instruction": "List main components and responsibilities in a table: name, responsibility, notes."},
            {"name": "APIs", "type": "table", "instruction": "List external/internal APIs: endpoint, purpose, auth."},
            {"name": "Data Model", "type": "text", "instruction": "Describe key data models and storage."},
            {"name": "Deployment", "type": "text", "instruction": "Describe deployment topology and CI/CD."},
            {"name": "Security", "type": "text", "instruction": "High level security considerations."},
            {"name": "Appendix", "type": "text", "instruction": "Optional appendices or references."}
        ]
        with open(PROMPT_FILE, "w", encoding="utf-8") as f:
            json.dump(default, f, indent=2)

    # create a simple template docx if not present
    if not os.path.exists(TEMPLATE_FILE):
        logger.info("Creating default HLD template at %s", TEMPLATE_FILE)
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        with open(PROMPT_FILE, 'r', encoding='utf-8') as f:
            sections = json.load(f)
        doc.add_heading('High Level Design', level=1)
        for s in sections:
            doc.add_heading(s['name'], level=2)
            doc.add_paragraph(s.get('instruction', ''))
        doc.save(TEMPLATE_FILE)
        logger.debug("Template created")


def scan_folder(path, max_files=200, max_bytes_per_file=8_192):
    files = []
    logger.debug("Scanning folder %s", path)
    binary_exts = {'.pyc', '.class', '.so', '.exe', '.dll', '.bin', '.jpg', '.jpeg', '.png', '.gif', '.zip', '.tar', '.gz'}
    for root, dirs, filenames in os.walk(path):
        # skip common large or irrelevant folders
        if any(part in ('__pycache__', '.git', 'node_modules') for part in root.split(os.sep)):
            logger.debug("Skipping directory in scan: %s", root)
            continue
        for fname in filenames:
            fpath = os.path.join(root, fname)
            # skip binary-like files by extension
            _, ext = os.path.splitext(fname.lower())
            if ext in binary_exts:
                logger.debug("Skipping binary file: %s", fpath)
                continue
            try:
                stat = os.path.getsize(fpath)
                if stat > 200_000:
                    logger.debug("Skipping large file: %s (%d bytes)", fpath, stat)
                    continue
                with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                    snippet = f.read(max_bytes_per_file)
                # only include files that have at least some readable text
                if not snippet.strip():
                    logger.debug("Skipping empty-text file: %s", fpath)
                    continue
                files.append({'path': os.path.relpath(fpath, path), 'size': stat, 'snippet': snippet})
            except Exception:
                logger.debug("Skipping file due to read error: %s", fpath)
                continue
            if len(files) >= max_files:
                break
    logger.info("Scanned %d files from %s", len(files), path)
    return files


def build_prompts(sections, files):
    files_summary = "\n".join([f"{f['path']}" for f in files])
    prompts = {}
    for s in sections:
        header = s['name']
        instr = s.get('instruction', '')
        prompt = f"{header}\nInstruction: {instr}\nFILES_SUMMARY:\n{files_summary}\n\nPlease generate the content for the section named '{header}'. Use tables if requested."
        prompts[header] = prompt
    return prompts


def render_docx(sections, outputs, out_path):
    doc = Document()
    doc.add_heading('High Level Design', level=1)
    for s in sections:
        name = s['name']
        typ = s.get('type', 'text')
        doc.add_heading(name, level=2)
        content = outputs.get(name, '')
        # If section is table type and content looks like lines with | separators, try building table
        if typ == 'table' and ('|' in content):
            lines = [l.strip() for l in content.splitlines() if l.strip()]
            # parse first line as header if it contains |
            rows = [l.split('|') for l in lines]
            rows = [[c.strip() for c in r if c is not None] for r in rows]
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(rows[0]):
                    hdr_cells[i].text = h
                for r in rows[1:]:
                    row_cells = table.add_row().cells
                    for i, c in enumerate(r):
                        if i < len(row_cells):
                            row_cells[i].text = c
        else:
            for para in content.splitlines():
                doc.add_paragraph(para)
    doc.save(out_path)


async def generate_from_path(path):
    logger.info("Generating HLD for path: %s", path)
    try:
        with open(PROMPT_FILE, 'r', encoding='utf-8') as f:
            sections = json.load(f)
    except Exception as e:
        logger.exception("Failed to read prompt file: %s", e)
        raise

    files = scan_folder(path)
    prompts = build_prompts(sections, files)
    outputs = {}

    for header, prompt in prompts.items():
        logger.debug("Generating section: %s (prompt length=%d)", header, len(prompt))
        try:
            text = await generate_text(prompt)
            outputs[header] = text
            logger.debug("Generated section '%s' (chars=%d)", header, len(text))
        except Exception as e:
            logger.exception("LLM generation failed for section %s: %s", header, e)
            outputs[header] = f"[Generation failed: {e}]"

    timestamp = int(time.time())
    out_filename = f"generated_hld_{timestamp}.docx"
    out_path = os.path.join('/tmp', out_filename)
    try:
        render_docx(sections, outputs, out_path)
        logger.info("Saved generated document to %s", out_path)
    except Exception as e:
        logger.exception("Failed to render docx: %s", e)
        raise
    return out_path
